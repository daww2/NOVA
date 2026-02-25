"""
Document Processor for RAG Pipeline.

FOCUS: PDF, DOCX, TXT, CSV processing
MUST: Extract text, preserve structure

Uses:
- pdfplumber: PDF text + table extraction
- python-docx: Word documents
- pandas: CSV/Excel
"""

import logging
from dataclasses import dataclass
from pathlib import Path
import pandas as pd

import pdfplumber
from docx import Document
from bs4 import BeautifulSoup   # Read HTML like a tree instead of raw text

logger = logging.getLogger(__name__)

@dataclass
class ProcessedDocument:
    """Extracted document content."""
    content: str
    metadata: dict
    source : str
    page_count : int = 1
    

class DocumentProcessor:
    """
    Process documents for RAG ingestion.
    
    Usage:
        processor = DocumentProcessor()
        doc = processor.process("document.pdf")
        chunks = chunker.chunk(doc.content)
    """
    SUPPORTED_TYPES = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".txt": "text",
        ".md": "text",
        ".csv": "csv",
        ".xlsx": "excel",
        ".html": "html",
    }
    
    def process(self, file_path: str) -> ProcessedDocument:
        """
        Process a document file.
        
        Args:
            file_path: Path to document
            
        Returns:
            ProcessedDocument with extracted content
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = path.suffix.lower() # [.pdf , .md, .csv , etc]
        doc_type = self.SUPPORTED_TYPES.get(ext)
        
        if not doc_type:
            raise ValueError(f"Unsupported file type: {ext}")


        # Route to appropriate extractor
        extractors = {
            "pdf": self._extract_pdf,
            "docx": self._extract_docx,
            "text": self._extract_text,
            "csv": self._extract_csv,
            "excel": self._extract_excel,
            "html": self._extract_html,
        }
        
        content, page_count  = extractors[doc_type](path)
        
        return ProcessedDocument(
            content=content,
            metadata={
                "filename": path.name,
                "file_type": ext,
                "file_size": path.stat().st_size,
            },
            source=str(path),
            page_count=page_count,
        )
        
    def _extract_pdf(self, path: Path) -> tuple[str, int]:
        """Extract text and tables from PDF using pdfplumber."""

        pages = []

        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                page_parts: list[str] = []

                # Extract tables first so we can exclude their bboxes from text
                tables = page.find_tables()
                table_bboxes = [t.bbox for t in tables]

                # Extract text outside tables
                if table_bboxes:
                    # Crop page to exclude table regions and extract remaining text
                    filtered = page.filter(
                        lambda obj: not any(
                            bbox[0] <= obj.get("x0", 0) <= bbox[2]
                            and bbox[1] <= obj.get("top", 0) <= bbox[3]
                            for bbox in table_bboxes
                        )
                    )
                    text = filtered.extract_text() or ""
                else:
                    text = page.extract_text() or ""

                if text.strip():
                    page_parts.append(text.strip())

                # Extract each table as a markdown-style table
                for table in tables:
                    rows = table.extract()
                    if not rows:
                        continue
                    md_rows = []
                    for i, row in enumerate(rows):
                        cells = [str(c).strip() if c else "" for c in row]
                        md_rows.append("| " + " | ".join(cells) + " |")
                        if i == 0:
                            md_rows.append("| " + " | ".join("---" for _ in cells) + " |")
                    page_parts.append("\n".join(md_rows))

                pages.append("\n\n".join(page_parts))

            page_count = len(pdf.pages)

        return "\f".join(pages), page_count
    
    def _extract_docx(self, path: Path) -> tuple[str, int]:
        """Extract text from Word document."""

        doc = Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:  # Fixed: table.rows (not table.row)
                row_text = " | ".join(cell.text for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs), 1

        
    def _extract_text(self, path: Path) -> tuple[str, int]:
        """Extract plain text."""
        content = path.read_text(encoding="utf-8", errors="ignore")
        return content, 1
    
    def _extract_csv(self, path: Path) -> tuple[str, int]:
        """Extract CSV as text."""
        
        df = pd.read_csv(path)
        # Convert to readable text format
        content = df.to_string(index=False)
        return content, 1
    
    def _extract_excel(self, path: Path) -> tuple[str, int]:
        """Extract Excel as text."""
                
        # Read all sheets
        sheets = pd.read_excel(path, sheet_name=None)
        
        content_parts = []
        for sheet_name, df in sheets.items():
            content_parts.append(f"## {sheet_name}\n{df.to_string(index=False)}")
        
        return "\n\n".join(content_parts), len(sheets)
    
    def _extract_html(self, path: Path) -> tuple[str, int]:
        """Extract text from HTML."""
        
        html = path.read_text(encoding="utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        
        # Remove scripts and style elements 
        for element in soup(["script", "style"]):
            element.decompose()
            
        text = soup.get_text(separator="\n")
        # Clean up whitespace
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        
        return "\n".join(lines), 1
    
    def process_batch(
        self,
        file_paths: list[str],
        skip_errors: bool = True,
    ) -> list[ProcessedDocument]:
        """Process multiple documents.""" 
        results = []
        
        for path in file_paths:
            try:
                doc = self.process(path)
                results.append(doc)
            except Exception as e:
                logger.error(f"Failed to process {path}: {e}")
                if not skip_errors:
                    raise
        
        return results